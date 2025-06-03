
# VERSION v16: GPT-3.5 исправляет слова, GPT-4 только на low probability. Вопросы — через правило.

USE_CORRECTOR = True  # Отключить, если не нужны правки слов


import argparse

# === Аргументы командной строки ===
parser = argparse.ArgumentParser(
    description="""
Пример использования:

  python3 transcribe_with_diarization_v20.py path/to/file.mp3 --lang en --gpt gpt-4 --probability-threshold 0.95

По умолчанию используется язык: ru, модель: gpt-3.5-turbo, порог уверенности: 0.55
""",
    formatter_class=argparse.RawDescriptionHelpFormatter
)

parser.add_argument("audio_file", help="Путь к .mp3 или .wav файлу")
parser.add_argument("--lang", choices=["ru", "en"], default="ru", help="Язык транскрипции (ru/en)")
parser.add_argument("--gpt", choices=["gpt-3.5-turbo", "gpt-4"], default="gpt-3.5-turbo", help="Модель GPT")
parser.add_argument("--probability-threshold", type=float, default=0.55,
                    help="Порог уверенности для отправки в GPT-4 (по умолчанию: 0.55)")
parser.add_argument("--skip-whisper", action="store_true", help="Пропустить этап Whisper, если есть сохранённый JSON")

args = parser.parse_args()

PROBABILITY_THRESHOLD = args.probability_threshold

audio_file = args.audio_file
LANGUAGE = args.lang
GPT_MODEL = args.gpt

import openai
system_prompt_ru = "Ты — корректировщик текста. Исправь ошибки распознавания, сохрани стиль, не сокращай, не убирай предлоги. Не исправляй сленг (например, 'косты', 'синк', 'бабки'). Ответь списком JSON объектов с полями 'speaker' и 'text'."
system_prompt_en = "You are a text corrector. Fix recognition errors, preserve the style, do not shorten or remove prepositions. Do not correct slang (e.g., 'costs', 'sync', 'bucks'). Return a list of JSON objects with fields 'speaker' and 'text'."
import json

def correct_text_with_gpt(text, speaker_label, use_gpt4=False):
    model = "gpt-4" if use_gpt4 else GPT_MODEL
    try:
        response = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt_ru if LANGUAGE == "ru" else system_prompt_en},
                {"role": "user", "content": f'{{"speaker": "{speaker_label}", "text": "{text}"}}'}
            ],
            max_tokens=500
        )
        cleaned = response.choices[0].message.content.strip()
        data = json.loads(cleaned)
        return data.get("speaker", speaker_label), data.get("text", text)
    except Exception as e:
        print(f"⚠️ GPT error: {e}")
        return speaker_label, text


# MASTER SCRIPT v7: diarize + transcribe with roles using Whisper (medium)


import ssl
ssl._create_default_https_context = ssl._create_unverified_context

import whisper
from resemblyzer import preprocess_wav, VoiceEncoder
from resemblyzer.hparams import sampling_rate
from sklearn.cluster import KMeans
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import timedelta, datetime
import numpy as np
import os, sys, json
import openai
system_prompt_ru = "Ты — корректировщик текста. Исправь ошибки распознавания, сохрани стиль, не сокращай, не убирай предлоги. Не исправляй сленг (например, 'косты', 'синк', 'бабки'). Ответь списком JSON объектов с полями 'speaker' и 'text'."
system_prompt_en = "You are a text corrector. Fix recognition errors, preserve the style, do not shorten or remove prepositions. Do not correct slang (e.g., 'costs', 'sync', 'bucks'). Return a list of JSON objects with fields 'speaker' and 'text'."

# === GPT Pricing model ===
pricing = {
    "gpt-3.5-turbo": {"prompt": 0.0005, "completion": 0.0015},
    "gpt-4":         {"prompt": 0.03,   "completion": 0.06}
}


# === Step 1: Get audio file ===
if len(sys.argv) > 1:
    audio_file = sys.argv[1]
else:
    audio_file = input("Введите путь к .mp3 или .wav файлу: ").strip()

if not os.path.exists(audio_file):
    print("❌ Файл не найден.")
    sys.exit(1)

# === Load OpenAI API key ===
api_key_path = os.path.join(os.path.dirname(audio_file), "openai_api_key.txt")
if not os.path.exists(api_key_path):
    api_key_path = os.path.expanduser("~/.openai_api_key")
if os.path.exists(api_key_path):
    with open(api_key_path, "r") as f:
        openai.api_key = f.read().strip()
else:
    print("❌ Файл с API ключом OpenAI не найден.")
    sys.exit(1)

# === Step 2: Diarization ===
print("🔊 Выполняем диаризацию...")
wav = preprocess_wav(audio_file)
encoder = VoiceEncoder()
_, embeds, wav_splits = encoder.embed_utterance(wav, return_partials=True, rate=16)

kmeans = KMeans(n_clusters=2, random_state=0).fit(embeds)
labels = kmeans.labels_

intervals = []
for i, label in enumerate(labels):
    start = wav_splits[i].start
    end = wav_splits[i].stop if i + 1 < len(wav_splits) else len(wav)
    intervals.append({
        "start": round(start / sampling_rate, 2),
        "end": round(end / sampling_rate, 2),
        "speaker": int(label)
    })

# === Save diarization to JSON ===
base_name = os.path.splitext(os.path.basename(audio_file))[0]
timestamp = datetime.now().strftime("%Y%m%d_%H%M")
json_path = f"{base_name}_diarization_{timestamp}.json"
with open(json_path, "w") as f:
    json.dump(intervals, f, indent=2)
print(f"✅ Диаризация завершена и сохранена в {json_path}")


whisper_result_path = f"{os.path.splitext(audio_file)[0]}_whisper.json"

# === Step 3: Whisper Transcription ===
if args.skip_whisper and os.path.exists(whisper_result_path):
    print("⏩ Пропускаем Whisper, используем сохранённый результат...")
    with open(whisper_result_path, "r", encoding="utf-8") as f:
        result = json.load(f)
else:
    print("🧠 Распознаём речь (Whisper medium)...")
    model = whisper.load_model("medium")
    result = model.transcribe(audio_file, language=LANGUAGE, word_timestamps=True)
    with open(whisper_result_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)


# === Step 4: Match speaker ===
def get_best_matching_speaker(start_time, end_time):
    durations = {0: 0.0, 1: 0.0}
    for entry in intervals:
        overlap_start = max(start_time, entry["start"])
        overlap_end = min(end_time, entry["end"])
        overlap = max(0.0, overlap_end - overlap_start)
        durations[entry["speaker"]] += overlap
    return 0 if durations[0] > durations[1] else 1


with open(whisper_result_path, "w", encoding="utf-8") as f:
    json.dump(result, f, ensure_ascii=False, indent=2)


if os.path.exists(whisper_result_path):
    with open(whisper_result_path, "r", encoding="utf-8") as f:
        result = json.load(f)
else:
    result = model.transcribe(...)


# === GPT-3.5 Question Enhancer ===
gpt35_requests = 0
gpt35_tokens_prompt = 0
gpt35_tokens_completion = 0
gpt4_requests = 0
gpt4_tokens_prompt = 0
gpt4_tokens_completion = 0
correction_log = []

# def enhance_questions_with_gpt(text, speaker_label):
#     global gpt35_requests, gpt35_tokens_prompt, gpt35_tokens_completion, correction_log
#     try:
#         payload = json.dumps({"speaker": speaker_label, "text": text}, ensure_ascii=False)
#         response = openai.ChatCompletion.create(
#             model="gpt-3.5-turbo",
#             messages=[
#                 {"role": "system", "content": """Ты редактор. Добавь в текст пропущенные знаки вопроса, если предложения начинаются с одного из следующих слов: что, как, где, почему, когда, зачем, кто, какой, какая, какие, каков, каково, откуда, сколько, чем, чего, чему, чьё, чья, чьи, чей. Также добавляй знак вопроса, если предложение начинается с фраз: ну как, ну что, ну где, ну почему, а как, а что, а где, а почему, а зачем, а кто и т.п. Также добавляй знак вопроса, если встречается частица 'ли' после подлежащего. Не изменяй стиль, не сокращай. Сохрани сленг и разговорные выражения. Верни JSON с полями: 'speaker' и 'text'."""},
#                 {"role": "user", "content": payload}
#             ],
#             max_tokens=500
#         )
#         usage = response.usage
#         gpt35_requests += 1
#         gpt35_tokens_prompt += usage.prompt_tokens
#         gpt35_tokens_completion += usage.completion_tokens
#         cleaned = json.loads(response.choices[0].message.content.strip())
#         correction_log.append({
#             "original": {"speaker": speaker_label, "text": text},
#             "corrected": cleaned
#         })
#         return cleaned.get("speaker", speaker_label), cleaned.get("text", text)
#     except Exception as e:
#         print(f"⚠️ GPT ошибка: {e}")
#         return speaker_label, text


def enhance_questions_with_gpt(text, speaker=None):
    question_words = {
        "что", "как", "где", "почему", "зачем", "когда", "какой", "какая", "какие", "каков", "откуда",
        "сколько", "чего", "чем", "чье", "чья", "чьи", "кому", "кого", "куда", "почем", "по-чему",
        "ну что", "ну как", "ну где", "а что", "а как", "а где", "а почему", "а зачем", "а когда"
    }

    cleaned_text = text.strip()
    text_lower = cleaned_text.lower()

    if not text_lower.endswith("?"):
        for w in question_words:
            if text_lower.startswith(w + " ") or text_lower.startswith(w + ","):
                cleaned_text += "?"
                break

    return speaker, cleaned_text


# === Step 5: Build Word Document ===
print("📄 Формируем Word-документ...")
doc = Document()
doc.add_heading(f"Транскрипция с ролями – {audio_file}", level=1)


for seg in result["segments"]:
    
    if not seg.get("text", "").strip():
        continue

    start = round(seg["start"], 2)
    end = round(seg["end"], 2)
    speaker_id = get_best_matching_speaker(start, end)
    speaker = "Интервьюер" if speaker_id == 0 else "Кандидат"

    timecode = str(timedelta(seconds=int(start)))
    doc.add_paragraph(f"[{timecode}]")

    # Проверка probability

    has_low_probability = False
    words = seg.get("words", [])

    if words:
        low_conf = [w for w in words if w.get("probability", 1.0) < PROBABILITY_THRESHOLD]
        has_low_probability = len(low_conf) > 0

    else:
        print(f"⚠️ Нет word-level probability в сегменте: '{seg['text']}'")
        # Альтернативная эвристика — по avg_logprob
        if seg.get("avg_logprob", 0) < -1.0:
            has_low_probability = True


    if has_low_probability:
        try:
            payload = json.dumps({"speaker": speaker, "text": seg["text"].strip()}, ensure_ascii=False)
            model_to_use = GPT_MODEL  # Используем только заданную модель

            response = openai.ChatCompletion.create(
                model=model_to_use,
                messages=[
                    {
                        "role": "system",
                        "content": "Ты редактор. Исправь ошибки в распознанном тексте. Не меняй стиль, не сокращай и не заменяй сленг. Верни JSON с полями: 'speaker' и 'text'."
                    },
                    {"role": "user", "content": payload}
                ],
                max_tokens=500
            )

            usage = response.usage
            if model_to_use == "gpt-4":
                gpt4_requests += 1
                gpt4_tokens_prompt += usage.prompt_tokens
                gpt4_tokens_completion += usage.completion_tokens
            else:
                gpt35_requests += 1
                gpt35_tokens_prompt += usage.prompt_tokens
                gpt35_tokens_completion += usage.completion_tokens

            cleaned = json.loads(response.choices[0].message.content.strip())
            final_speaker = cleaned.get("speaker", speaker)
            enhanced_text = cleaned.get("text", seg["text"].strip())
            correction_log.append({
                "original": {"speaker": speaker, "text": seg["text"].strip()},
                "corrected": cleaned,
                "used_model": model_to_use
            })
        except Exception as e:
            print(f"⚠️ GPT ошибка: {e}")
            final_speaker, enhanced_text = enhance_questions_with_gpt(seg["text"], speaker)
    else:
        final_speaker, enhanced_text = enhance_questions_with_gpt(seg["text"], speaker)


    p = doc.add_paragraph()
    run = p.add_run(f"{final_speaker}: ")
    run.bold = True
    run.font.size = Pt(11)
    run_text = p.add_run(enhanced_text)
    run_text.font.size = Pt(11)

    if final_speaker == "Интервьюер":
        run_text.font.color.rgb = RGBColor(0, 128, 0)
    if enhanced_text.strip().endswith("?"):
        run_text.font.color.rgb = RGBColor(128, 0, 128)


# === Save output files ===
doc_timestamp = datetime.now().strftime("%Y%m%d_%H%M")
output_file = f"{os.path.splitext(audio_file)[0]}_with_roles_{doc_timestamp}.docx"
doc.save(output_file)
print(f"✅ Готово: {output_file}")

# === Save GPT stats ===
stats_path = os.path.splitext(audio_file)[0] + f"_gpt_stats_{doc_timestamp}.txt"
with open(stats_path, "w") as f:
    f.write(f"🎯 Порог probability (probability): {PROBABILITY_THRESHOLD}\n")
    f.write(f"💬 Запросов GPT-3.5: {gpt35_requests}\n")
    f.write(f"🔢 Prompt токенов GPT-3.5: {gpt35_tokens_prompt}\n")
    f.write(f"🔢 Completion токенов GPT-3.5: {gpt35_tokens_completion}\n")
    f.write(f"💬 Запросов GPT-4: {gpt4_requests}\n")
    f.write(f"🔢 Prompt токенов GPT-4: {gpt4_tokens_prompt}\n")
    f.write(f"🔢 Completion токенов GPT-4: {gpt4_tokens_completion}\n")
    cost_35 = (gpt35_tokens_prompt * pricing['gpt-3.5-turbo']['prompt'] + gpt35_tokens_completion * pricing['gpt-3.5-turbo']['completion']) / 1000
    cost_4 = (gpt4_tokens_prompt * pricing['gpt-4']['prompt'] + gpt4_tokens_completion * pricing['gpt-4']['completion']) / 1000
    total = cost_35 + cost_4
    f.write(f"💰 Оценка стоимости:\n")
    f.write(f"GPT-3.5: ~$ {cost_35:.4f}\n")
    f.write(f"GPT-4:   ~$ {cost_4:.4f}\n")
    f.write(f"ИТОГО:  ~$ {total:.4f}\n")
print(f"📊 Статистика GPT сохранена в {stats_path}")

# === Save corrections log ===
log_path = os.path.splitext(audio_file)[0] + f"_gpt_corrections_log_{doc_timestamp}.json"
with open(log_path, "w", encoding="utf-8") as f:
    json.dump(correction_log, f, indent=2, ensure_ascii=False)
print(f"📁 Лог правок сохранён в {log_path}")



# Пример вызова внутри основного цикла обработки сегментов:
if USE_CORRECTOR:
    new_speaker, cleaned_text = correct_text_with_gpt(seg["text"].strip(), speaker, has_low_probability)
else:
    new_speaker, cleaned_text = speaker, seg["text"].strip()
